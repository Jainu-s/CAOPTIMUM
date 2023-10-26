import openpyxl
import re

#
# def extract_data_from_text(file_name):
#     with open(file_name, 'r') as file:
#         lines = file.readlines()
#
#     columns = []
#     all_rows = []
#     current_row = {}
#     for line in lines:
#         line = line.strip()
#         if line.endswith(':'):
#             key = line[:-1]
#             if key in current_row:
#                 all_rows.append(current_row)
#                 current_row = {}
#             if key not in columns:
#                 columns.append(key)
#         elif line and key:
#             current_row[key] = line
#
#     if current_row:
#         all_rows.append(current_row)
#
#     return columns, all_rows
#
#
#
# def action_exists(text, actions):
#     for action in actions:
#         if action in text:
#             return True
#     return False
#
# def format_steps_content(content):
#     actions = ['click', 'insert', 'hover', 'url', 'browser','select']
#     delimiters = ['.', ',', 'and']
#
#     # Splitting content based on delimiters
#     segments = re.split('|'.join(map(re.escape, delimiters)), content)
#
#     # Only keep segments with at least one action
#     segments = [segment.strip() for segment in segments if action_exists(segment, actions)]
#
#     # Join segments with a newline and a prefix number
#     formatted_content = "\n".join(f"{i+1}. {segment}" for i, segment in enumerate(segments))
#
#     return formatted_content
#
# def write_to_excel(columns, all_rows, excel_file):
#     wb = openpyxl.Workbook()
#     ws = wb.active
#
#     # Write column headers
#     for col_num, col_name in enumerate(columns, 1):
#         ws.cell(row=1, column=col_num).value = col_name
#
#     # Write row data
#     for row_num, row_data in enumerate(all_rows, 2):
#         for col_num, col_name in enumerate(columns, 1):
#             value = row_data.get(col_name, '')
#             if "Step" in col_name or "Steps" in col_name:
#                 value = format_steps_content(value)
#             ws.cell(row=row_num, column=col_num).value = value
#
#     wb.save(excel_file)
#
#
# if __name__ == '__main__':
#     columns, all_rows = extract_data_from_text(r"C:\Users\abdul\PycharmProjects\Automation_Optimum\Testcases\swiggy.txt")
#     write_to_excel(columns, all_rows, "output.xlsx")


# import re
# import pandas as pd
# from docx import Document
#
#
# def extract_data_from_text(file_name):
#     with open(file_name, 'r') as file:
#         lines = file.readlines()
#     return extract_data_from_lines(lines)
#
#
# def extract_data_from_word(file_name):
#     doc = Document(file_name)
#     lines = [p.text for p in doc.paragraphs if p.text.strip() != '']
#     return extract_data_from_lines(lines)
#
#
# def extract_data_from_lines(lines):
#     columns = []
#     all_rows = []
#     current_row = {}
#     for line in lines:
#         line = line.strip()
#         if line.endswith(':'):
#             key = line[:-1]
#             if key in current_row:
#                 all_rows.append(current_row)
#                 current_row = {}
#             if key not in columns:
#                 columns.append(key)
#         elif line:
#             current_row[key] = line
#
#     if current_row:
#         all_rows.append(current_row)
#
#     return columns, all_rows
#
#
# def action_exists(text, actions):
#     for action in actions:
#         if action in text:
#             return True
#     return False
#
#
# def format_steps_content(content):
#     actions = ['click', 'insert', 'hover', 'url', 'browser', 'select','enter']
#     delimiters = ['.', ',', 'and']
#     segments = re.split('|'.join(map(re.escape, delimiters)), content)
#     segments = [segment.strip() for segment in segments if action_exists(segment, actions)]
#     formatted_content = "\n".join(f"{i + 1}. {segment}" for i, segment in enumerate(segments))
#     return formatted_content
#
#
# def create_dataframe(columns, all_rows):
#     # Convert the columns and rows to a pandas DataFrame
#     df = pd.DataFrame(all_rows, columns=columns)
#
#     # If you have columns with step-based content, format them:
#     for col_name in df.columns:
#         if "Step" in col_name or "Steps" in col_name:
#             df[col_name] = df[col_name].apply(format_steps_content)
#
#     return df
#
# def save_to_excel(df, file_name):
#     """Save a pandas DataFrame to an Excel file."""
#     with pd.ExcelWriter(file_name, engine='openpyxl') as writer:
#         df.to_excel(writer, index=False)
#
# def text_to_excel_main(file_path=None):
#     file_path = r"C:\Users\abdul\PycharmProjects\Automation_Optimum\Testcases\swiggy.docx"
#     if file_path.endswith('.docx'):
#         columns, all_rows = extract_data_from_word(file_path)
#     else:
#         columns, all_rows = extract_data_from_text(file_path)
#
#     if columns and all_rows:
#         df = create_dataframe(columns, all_rows)
#         save_to_excel(df, 'text_doc_excel.xlsx')
#
# text_to_excel_main()




import re
import pandas as pd
from docx import Document

def extract_data_from_text(file_name):
    with open(file_name, 'r') as file:
        lines = file.readlines()
    return extract_data_from_lines(lines)

def extract_data_from_word(file_name):
    doc = Document(file_name)
    lines = [p.text for p in doc.paragraphs if p.text.strip() != '']
    return extract_data_from_lines(lines)

def extract_data_from_lines(lines):
    columns = []
    original_columns = []  # Store the original case of the columns here
    all_rows = []
    current_row = {}
    for line in lines:
        line = line.strip()
        if line.lower().endswith(':'):
            original_key = line[:-1]  # Retain original key
            key = original_key.lower()  # Convert to lowercase for internal operations
            if key in current_row:
                all_rows.append(current_row)
                current_row = {}
            if key not in columns:
                columns.append(key)
                original_columns.append(original_key)  # Store the original column name
        elif line:
            current_row[key] = line

    if current_row:
        all_rows.append(current_row)

    return columns, original_columns, all_rows  # Return the original column names as well

def action_exists(text, actions):
    for action in actions:
        if action in text:
            return True
    return False

def format_steps_content(content):
    actions = ['click', 'insert', 'hover', 'url', 'browser', 'select','enter']
    delimiters = ['.', ',', 'and']
    segments = re.split('|'.join(map(re.escape, delimiters)), content)
    segments = [segment.strip() for segment in segments if action_exists(segment, actions)]
    formatted_content = "\n".join(f"{i + 1}. {segment}" for i, segment in enumerate(segments))
    return formatted_content

def create_dataframe(columns, original_columns, all_rows):
    # Convert the columns and rows to a pandas DataFrame
    df = pd.DataFrame(all_rows, columns=columns)

    # If you have columns with step-based content, format them:
    for col_name in df.columns:
        if "step" in col_name.lower():
            df[col_name] = df[col_name].apply(format_steps_content)

    # Convert columns back to their original case before output
    df.columns = original_columns

    return df

def save_to_excel(df, file_name):
    """Save a pandas DataFrame to an Excel file."""
    with pd.ExcelWriter(file_name, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)

def text_to_excel_main(file_path=None):
    file_path = r"C:\Users\Optimum.LAPTOP-SQLU1RCT\PycharmProjects\FAP\CAOPTIMUM\Testcases\swiggy.docx"
    if file_path.endswith('.docx'):
        columns, original_columns, all_rows = extract_data_from_word(file_path)
    else:
        columns, original_columns, all_rows = extract_data_from_text(file_path)

    if columns and all_rows:
        df = create_dataframe(columns, original_columns, all_rows)
        save_to_excel(df, 'text_doc_excel.xlsx')

text_to_excel_main()
